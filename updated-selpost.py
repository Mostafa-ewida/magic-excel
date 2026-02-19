from flask import Flask, request, send_file, render_template
import docx
from openpyxl import Workbook, load_workbook
import os
import pandas as pd
import logging
from pathlib import Path
from typing import Generator, Union, Optional, Dict, Any



class LargeExcelReader:
    def __init__(self, chunk_size: int = 10000):
        self.chunk_size = chunk_size
        self.logger = self._setup_logger()

    def _setup_logger(self) -> logging.Logger:
        logger = logging.getLogger('LargeExcelReader')
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
            logger.setLevel(logging.INFO)
        return logger

    def read_excel(
    self,
    file_path: Union[str, Path],
    sheet_name: Optional[Union[str, int]] = 0,
    headers: Optional[Union[int, list]] = 0,
    usecols: Optional[Union[list, str]] = None,
    dtype: Optional[Dict] = None
    ) -> Generator[pd.DataFrame, None, None]:
        """
        Read a large Excel file in chunks with enhanced error recovery and partial data retrieval.
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            file_size = file_path.stat().st_size / (1024 * 1024)
            self.logger.info(f"Processing file: {file_path.name} (Size: {file_size:.2f} MB)")

            # Use ExcelFile to avoid reopening the file multiple times
            with pd.ExcelFile(file_path, engine='openpyxl') as excel_file:
                # Get the sheet name if it's a number
                if isinstance(sheet_name, int):
                    available_sheets = excel_file.sheet_names
                    if sheet_name >= len(available_sheets):
                        raise ValueError(f"Sheet index {sheet_name} is out of bounds")
                    sheet_name = available_sheets[sheet_name]

                # Get total number of rows in the sheet
                sheet = excel_file.book[sheet_name]
                total_rows = sheet.max_row - (headers + 1 if headers is not None else 0)
                
                # Calculate number of chunks
                num_chunks = (total_rows + self.chunk_size - 1) // self.chunk_size
                
                # Initialize columns
                self.columns = None
                successful_rows = 0
                
                # Read the file in chunks with advanced error recovery
                for chunk_number in range(num_chunks):
                    try:
                        start_row = chunk_number * self.chunk_size + (headers + 1 if headers is not None else 0)
                        end_row = min(start_row + self.chunk_size, total_rows + (headers + 1 if headers is not None else 0))
                        
                        # Read specific rows with comprehensive error handling
                        try:
                            chunk = pd.read_excel(
                                excel_file,
                                sheet_name=sheet_name,
                                header=headers if chunk_number == 0 else None,
                                usecols=usecols,
                                dtype=dtype,
                                skiprows=range(1, start_row) if start_row > 1 else None,
                                nrows=self.chunk_size
                            )
                        except Exception as read_error:
                            self.logger.warning(f"Primary read failed for chunk {chunk_number + 1}: {read_error}")
                            # Fallback strategy with reduced chunk size
                            try:
                                reduced_chunk_size = max(1000, self.chunk_size // 2)
                                chunk = pd.read_excel(
                                    excel_file,
                                    sheet_name=sheet_name,
                                    header=headers if chunk_number == 0 else None,
                                    usecols=usecols,
                                    dtype=dtype,
                                    skiprows=range(1, start_row) if start_row > 1 else None,
                                    nrows=reduced_chunk_size
                                )
                            except Exception as fallback_error:
                                self.logger.error(f"Fallback read failed for chunk {chunk_number + 1}: {fallback_error}")
                                continue
                        
                        # Handle columns consistency
                        if chunk_number == 0:
                            self.columns = chunk.columns
                        else:
                            chunk.columns = self.columns
                        
                        # Clean the data
                        chunk = chunk.replace({pd.NA: None, pd.NaT: None})
                        
                        # Update successful rows count
                        successful_rows += len(chunk)
                        
                        self.logger.info(f"Processing chunk {chunk_number + 1}/{num_chunks} - "
                                    f"Rows processed: {successful_rows:,}/{total_rows:,}")
                        
                        yield chunk

                    except Exception as chunk_error:
                        self.logger.warning(f"Error processing chunk {chunk_number + 1}: {str(chunk_error)}")
                        continue

                self.logger.info(f"Completed processing {successful_rows:,} rows out of {total_rows:,} total rows")

        except Exception as e:
            self.logger.error(f"Critical error reading Excel file: {str(e)}")
            raise

def concatenate_excel_sheets(files):
    MAX_ROWS_PER_SHEET = 1000000
    all_data = []
    reader = LargeExcelReader(chunk_size=50000)
    
    # Create temporary directory for saving files
    os.makedirs("temp", exist_ok=True)
    
    for file in files:
        try:
            # Save the uploaded file temporarily
            temp_file_path = os.path.join("temp", file.filename)
            file.save(temp_file_path)
            
            # Process the file in chunks with error recovery
            for chunk in reader.read_excel(temp_file_path):
                if not chunk.empty:
                    # Flatten all columns into a single list
                    values = chunk.values.flatten()
                    filtered_values = [v for v in values if pd.notna(v)]
                    all_data.extend(filtered_values)
                
            # Clean up temporary file
            os.remove(temp_file_path)
            
        except Exception as e:
            print(f"Error processing file {file.filename}: {e}")
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            continue

    # Save concatenated data to Excel file with multiple sheets
    output_path = os.path.join("temp", "concatenated.xlsx")
    
    if not all_data:
        # Handle case where no data was successfully processed
        wb = Workbook()
        ws = wb.active
        ws['A1'] = "No data could be processed from the input files"
        wb.save(output_path)
        return output_path

    try:
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Split data into chunks that fit within Excel's limits
            total_rows = len(all_data)
            num_sheets = (total_rows - 1) // MAX_ROWS_PER_SHEET + 1
            
            for sheet_num in range(num_sheets):
                start_idx = sheet_num * MAX_ROWS_PER_SHEET
                end_idx = min((sheet_num + 1) * MAX_ROWS_PER_SHEET, total_rows)
                
                chunk_data = all_data[start_idx:end_idx]
                df_chunk = pd.DataFrame(chunk_data, columns=["A"])
                
                sheet_name = f"Sheet_{sheet_num + 1}" if sheet_num > 0 else "Sheet1"
                df_chunk.to_excel(writer, sheet_name=sheet_name, index=False)
    
    except Exception as e:
        # Handle any errors during saving
        print(f"Error saving concatenated file: {e}")
        wb = Workbook()
        ws = wb.active
        ws['A1'] = f"Error saving concatenated data: {str(e)}"
        wb.save(output_path)

    return output_path

def search_keyword_in_first_column(file, keyword):
    result_wb = Workbook()
    result_ws = result_wb.active
    result_ws.append(["Column A"])  # Header
    
    try:
        # Initialize LargeExcelReader
        reader = LargeExcelReader(chunk_size=50000)
        
        # Flag to track if any matches are found
        matches_found = False
        
        # Process the file in chunks
        for chunk in reader.read_excel(file):
            if chunk.empty:
                continue
            
            # Get the first column name
            first_col = chunk.columns[0]
            
            # Search for keyword in the first column
            matches = chunk[chunk[first_col].astype(str).str.contains(keyword, case=False, na=False)]
            
            if not matches.empty:
                matches_found = True
                for _, row in matches.iterrows():
                    result_ws.append([str(row[first_col])])
        
        # If no matches were found, add a message
        if not matches_found:
            result_ws.append([f"No matches found for: {keyword}"])
    
    except Exception as e:
        print(f"Error searching file: {e}")
        result_ws.append([f"Error processing file: {str(e)}"])
    
    # Save results to a temporary file
    output_path = os.path.join("temp", f"search_results_{keyword}.xlsx")
    result_wb.save(output_path)
    return output_path
# Flask routes stay the same

app = Flask(__name__)

# Function to read lines from a Word document
def read_word_file(word_file_path):
    doc = docx.Document(word_file_path)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    return lines

# Function to write lines to an Excel file
def write_to_excel(lines, excel_file_path):
    wb = Workbook()
    ws = wb.active

    # Writing each line in the first column
    for idx, line in enumerate(lines, start=1):
        ws[f"A{idx}"] = line

    wb.save(excel_file_path)



@app.route('/')
def upload_form():
    return render_template('upload.html')

# Route to handle Word file upload and conversion
@app.route('/convert', methods=['POST'])
def convert_file():
    if 'word_file' not in request.files:
        return "No file part", 400

    file = request.files['word_file']

    if file.filename == '':
        return "No selected file", 400

    # Save the uploaded Word document temporarily
    word_file_path = os.path.join("temp", "uploaded.docx")
    file.save(word_file_path)

    # Read lines from Word document
    lines = read_word_file(word_file_path)

    # Define Excel file path
    excel_file_path = os.path.join("temp", "converted.xlsx")

    # Write lines to Excel file
    write_to_excel(lines, excel_file_path)

    # Remove the Word file after conversion
    os.remove(word_file_path)

    # Provide download link for the Excel file
    return send_file(excel_file_path, as_attachment=True, download_name="converted.xlsx")



# Route to handle file upload and concatenation
@app.route('/concatenate', methods=['POST'])
def concatenate_files():
    if 'excel_files' not in request.files:
        return "No files part", 400

    files = request.files.getlist('excel_files')

    if not files or files[0].filename == '':
        return "No files selected", 400

    try:
        # Pass the file objects directly to concatenate_excel_sheets
        concatenated_file_path = concatenate_excel_sheets(files)
        
        # Provide download link for the concatenated Excel file
        return send_file(
            concatenated_file_path,
            as_attachment=True,
            download_name="concatenated.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    except ValueError as ve:
        # Handle the case where no valid data was found
        return f"Error: {str(ve)}", 400
    except Exception as e:
        # Handle other errors
        return f"Error processing files: {str(e)}", 500

# Route to handle keyword search
@app.route('/search', methods=['POST'])
def search_keyword():
    if 'search_excel_file' not in request.files or 'search_keyword' not in request.form:
        return "File or keyword not provided", 400

    file = request.files['search_excel_file']
    keyword = request.form['search_keyword']

    if file.filename == '' or keyword == '':
        return "File or keyword not provided", 400

    # Save the uploaded file temporarily
    excel_file_path = os.path.join("temp", file.filename)
    file.save(excel_file_path)
    print("this is the file path", excel_file_path)
    # Search for the keyword in the Excel file
    result_file_path = search_keyword_in_first_column(excel_file_path, keyword)

    # Remove the original file
    os.remove(excel_file_path)

    # Provide download link for the result file
    return send_file(result_file_path, as_attachment=True, download_name=f"{keyword}.xlsx")



# Function to convert Excel to Word
def excel_to_word(excel_path, word_path):
    logger = logging.getLogger("ExcelToWord")

    if not logger.handlers:
        handler = logging.StreamHandler()
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        logger.addHandler(handler)
        logger.setLevel(logging.INFO)

    try:
        logger.info(f"Starting Excel → Word conversion: {excel_path}")

        reader = LargeExcelReader(chunk_size=50000)
        doc = docx.Document()

        total_rows = 0
        table_created = False

        for chunk in reader.read_excel(excel_path):
            if chunk.empty:
                continue

            # Create table only once
            if not table_created:
                table = doc.add_table(rows=1, cols=len(chunk.columns))

                # Add headers
                for col_idx, column in enumerate(chunk.columns):
                    table.rows[0].cells[col_idx].text = str(column)

                table_created = True

            # Append rows
            for _, row in chunk.iterrows():
                row_cells = table.add_row().cells
                for col_idx, cell in enumerate(row):
                    row_cells[col_idx].text = str(cell) if pd.notna(cell) else ""

                total_rows += 1

            logger.info(f"Processed {total_rows} rows so far...")

        if not table_created:
            doc.add_paragraph("No data found in Excel file.")

        doc.save(word_path)

        logger.info(f"Excel → Word conversion completed successfully. Total rows: {total_rows}")

    except Exception as e:
        logger.error(f"Excel → Word conversion failed: {str(e)}")
        raise


@app.route('/excel-to-word', methods=['POST'])
def convert_excel_to_word():
    if 'excel_file' not in request.files:
        return "No file uploaded", 400

    file = request.files['excel_file']

    if file.filename == '':
        return "No selected file", 400

    os.makedirs("temp", exist_ok=True)

    try:
        excel_path = os.path.join("temp", file.filename)
        file.save(excel_path)

        base_name = os.path.splitext(file.filename)[0]
        word_filename = f"{base_name}.docx"
        word_path = os.path.join("temp", word_filename)

        # Convert
        excel_to_word(excel_path, word_path)

        # Remove Excel safely
        if os.path.exists(excel_path):
            os.remove(excel_path)

        return send_file(
            word_path,
            as_attachment=True,
            download_name=word_filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        return f"Error converting file: {str(e)}", 500

if __name__ == '__main__':
    # Ensure the temp directory exists
    os.makedirs("temp", exist_ok=True)
    app.run(host='0.0.0.0', port=8080, debug=True)
